"""Gera documento Word do Modelo 1 a partir da planilha Excel formatada."""

import os
import shutil
from copy import deepcopy

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from core.formatador_sinapi.paths import caminho_modelo_word

TEMPLATE_PADRAO = "Modelo 1 - Word.docx"
ABA_ORCAMENTO = "Orçamento Formatado"

CABECALHO_ORCAMENTO = [
    "Item", "Código Sinapi", "Descrição", "Un.", "Qtd.",
    "Preço s/ BDI", "Preço c/ BDI", "Total s/ BDI", "Total c/ BDI",
]

LARGURAS_ORCAMENTO_CM = [1.25, 1.5, 5.25, 1.0, 1.25, 1.75, 1.75, 2.0, 3.5]
COL_ROTULO_TOTAL = 11
COL_VALOR_TOTAL = 12

FONTE_NOME = "Calibri"
FONTE_ITEM = Pt(10)
FONTE_SECAO = Pt(11)
FONTE_EXTENSO = Pt(9)
FONTE_REFERENCIA = "Arial"
FONTE_TOTAL_NORMAL = Pt(12)
FONTE_TOTAL_DESTAQUE = Pt(14)
FONTE_CONCLUSAO_TEXTO = Pt(12)
FONTE_CONCLUSAO_MOEDA = Pt(12)
FONTE_CONCLUSAO_EXTENSO = Pt(9)


def formatar_moeda(valor):
    if valor is None or valor == "":
        return ""
    if isinstance(valor, str):
        return valor.strip()
    return (
        f"R$ {valor:,.2f}"
        .replace(",", "X")
        .replace(".", ",")
        .replace("X", ".")
    )


def formatar_numero(valor):
    if valor is None or valor == "":
        return ""
    if isinstance(valor, (int, float)):
        if float(valor) == int(valor):
            return str(int(valor))
        texto = f"{valor:.4f}".rstrip("0").rstrip(".")
        return texto.replace(".", ",")
    return str(valor).strip()


def definir_texto_celula(celula, texto):
    texto = "" if texto is None else str(texto)
    if celula.paragraphs:
        paragrafo = celula.paragraphs[0]
        if paragrafo.runs:
            paragrafo.runs[0].text = texto
            for run in paragrafo.runs[1:]:
                run.text = ""
            return
        paragrafo.text = texto
        return
    celula.text = texto


def aplicar_fonte_celula(celula, tamanho, negrito=False, nome_fonte=FONTE_NOME):
    for paragrafo in celula.paragraphs:
        for run in paragrafo.runs:
            run.font.name = nome_fonte
            run.font.size = tamanho
            run.font.bold = negrito


def substituir_texto_celula(celula, texto, nome_fonte, tamanho, negrito=False, alinhamento=None):
    texto = "" if texto is None else str(texto)

    while len(celula.paragraphs) > 1:
        elemento = celula.paragraphs[-1]._element
        elemento.getparent().remove(elemento)

    paragrafo = celula.paragraphs[0] if celula.paragraphs else celula.add_paragraph()
    for filho in list(paragrafo._p):
        if filho.tag.endswith("}r"):
            paragrafo._p.remove(filho)

    run = paragrafo.add_run(texto)
    run.font.name = nome_fonte
    run.font.size = tamanho
    run.font.bold = negrito

    if alinhamento is not None:
        paragrafo.alignment = alinhamento


def configurar_run(run, texto, nome_fonte, tamanho, negrito=False):
    run.text = texto
    run.font.name = nome_fonte
    run.font.size = tamanho
    run.font.bold = negrito


def limpar_celula_conclusao(celula):
    while len(celula.paragraphs) > 1:
        elemento = celula.paragraphs[-1]._element
        elemento.getparent().remove(elemento)

    paragrafo = celula.paragraphs[0]
    for filho in list(paragrafo._p):
        if filho.tag.endswith("}r"):
            paragrafo._p.remove(filho)


def aplicar_larguras_colunas(tabela, larguras_cm):
    grade = tabela._tbl.tblGrid
    if grade is None:
        return
    for indice, largura in enumerate(larguras_cm):
        if indice < len(grade.gridCol_lst):
            grade.gridCol_lst[indice].w = int(Cm(largura))


def celulas_unicas(linha):
    vistas = set()
    resultado = []
    for indice, celula in enumerate(linha.cells):
        chave = id(celula._tc)
        if chave in vistas:
            continue
        vistas.add(chave)
        resultado.append((indice, celula))
    return resultado


def obter_cabecalho_tabela(tabela):
    if not tabela.rows:
        return []
    vistas = set()
    cabecalho = []
    for celula in tabela.rows[0].cells:
        chave = id(celula._tc)
        if chave in vistas:
            continue
        vistas.add(chave)
        cabecalho.append(celula.text.strip())
    return cabecalho


def limpar_linhas_tabela(tabela, manter_cabecalho=1):
    while len(tabela.rows) > manter_cabecalho:
        tabela._tbl.remove(tabela.rows[manter_cabecalho]._tr)


def adicionar_linha_clonada(tabela, tr_prototipo):
    tabela._tbl.append(deepcopy(tr_prototipo))
    return tabela.rows[-1]


def preencher_linha_secao(linha, item, descricao, extenso, total):
    valores = {0: item, 1: "", 2: descricao, 3: extenso, 8: formatar_moeda(total)}
    for indice, celula in celulas_unicas(linha):
        if indice in valores:
            definir_texto_celula(celula, valores[indice])
            if indice == 3:
                aplicar_fonte_celula(celula, FONTE_EXTENSO)
            else:
                aplicar_fonte_celula(celula, FONTE_SECAO)


def preencher_linha_item(linha, valores):
    for indice, celula in celulas_unicas(linha):
        if indice < len(valores):
            definir_texto_celula(celula, valores[indice])
            aplicar_fonte_celula(celula, FONTE_ITEM)


def preencher_linha_extenso(linha, texto):
    for indice, celula in celulas_unicas(linha):
        if indice == 3:
            definir_texto_celula(celula, texto)
            aplicar_fonte_celula(celula, FONTE_EXTENSO)
        elif indice not in (0,):
            definir_texto_celula(celula, "")


def preencher_tabela_totais(tabela, totais):
    for indice, (rotulo, valor) in enumerate(totais[: len(tabela.rows)]):
        linha = tabela.rows[indice]
        definir_texto_celula(linha.cells[0], rotulo)
        definir_texto_celula(linha.cells[1], f" {formatar_moeda(valor)} ")

        if indice < 2:
            aplicar_fonte_celula(linha.cells[0], FONTE_TOTAL_NORMAL, negrito=False)
            aplicar_fonte_celula(linha.cells[1], FONTE_TOTAL_NORMAL, negrito=False)
        else:
            aplicar_fonte_celula(linha.cells[0], FONTE_TOTAL_DESTAQUE, negrito=True)
            aplicar_fonte_celula(linha.cells[1], FONTE_TOTAL_DESTAQUE, negrito=True)


def reconstruir_tabela_orcamento(tabela, linhas):
    if len(tabela.rows) < 3:
        raise ValueError("A tabela de orçamento do template não possui linhas de protótipo.")

    proto_secao = deepcopy(tabela.rows[1]._tr)
    proto_item = deepcopy(tabela.rows[2]._tr)
    proto_extenso = deepcopy(tabela.rows[2]._tr)

    limpar_linhas_tabela(tabela)

    for linha_dados in linhas:
        tipo = linha_dados["tipo"]
        if tipo == "secao":
            linha = adicionar_linha_clonada(tabela, proto_secao)
            preencher_linha_secao(
                linha,
                linha_dados["item"],
                linha_dados["descricao"],
                linha_dados["extenso"],
                linha_dados["total"],
            )
        elif tipo == "item":
            linha = adicionar_linha_clonada(tabela, proto_item)
            preencher_linha_item(
                linha,
                [
                    linha_dados["item"],
                    linha_dados["codigo"],
                    linha_dados["descricao"],
                    linha_dados["un"],
                    formatar_numero(linha_dados["qtd"]),
                    formatar_moeda(linha_dados["preco_s_bdi"]),
                    formatar_moeda(linha_dados["preco_c_bdi"]),
                    formatar_moeda(linha_dados["total_s_bdi"]),
                    formatar_moeda(linha_dados["total"]),
                ],
            )
        elif tipo == "extenso":
            linha = adicionar_linha_clonada(tabela, proto_extenso)
            preencher_linha_extenso(linha, linha_dados["texto"])

    aplicar_larguras_colunas(tabela, LARGURAS_ORCAMENTO_CM)


def encontrar_tabelas_modelo1(documento):
    tabela_orcamento = None
    tabelas_totais = []

    for tabela in documento.tables:
        cabecalho = obter_cabecalho_tabela(tabela)
        if cabecalho == CABECALHO_ORCAMENTO:
            tabela_orcamento = tabela
        elif cabecalho and cabecalho[0] == "Total sem BDI":
            tabelas_totais.append(tabela)

    return tabela_orcamento, tabelas_totais


def ler_metadados_planilha(planilha):
    referencia_sinapi = ""
    texto_extenso_total = ""

    for numero_linha in range(planilha.max_row, 1, -1):
        valor = planilha.cell(numero_linha, COL_ROTULO_TOTAL).value
        if not valor:
            continue
        texto = str(valor).strip()
        if texto.upper().startswith("SINAPI") and not referencia_sinapi:
            referencia_sinapi = texto
        elif "reais" in texto.lower() and not texto_extenso_total:
            texto_extenso_total = texto

    return referencia_sinapi, texto_extenso_total


def ler_orcamento_formatado(planilha):
    linhas = []
    totais = []

    for numero_linha in range(2, planilha.max_row + 1):
        rotulo_total = planilha.cell(numero_linha, COL_ROTULO_TOTAL).value
        valor_total = planilha.cell(numero_linha, COL_VALOR_TOTAL).value

        if rotulo_total and str(rotulo_total).strip().startswith("Total"):
            totais.append((str(rotulo_total).strip(), valor_total))
            continue

        if rotulo_total and str(rotulo_total).strip().upper().startswith("SINAPI"):
            continue

        if (
            rotulo_total
            and valor_total is None
            and isinstance(rotulo_total, str)
            and "reais" in rotulo_total.lower()
        ):
            continue

        item = planilha.cell(numero_linha, 1).value
        codigo = planilha.cell(numero_linha, 2).value
        descricao = planilha.cell(numero_linha, 3).value
        coluna4 = planilha.cell(numero_linha, 4).value
        quantidade = planilha.cell(numero_linha, 5).value
        preco_s_bdi = planilha.cell(numero_linha, 6).value
        preco_c_bdi = planilha.cell(numero_linha, 7).value
        total_s_bdi = planilha.cell(numero_linha, 8).value
        total = planilha.cell(numero_linha, 9).value

        valores = [
            item, codigo, descricao, coluna4, quantidade,
            preco_s_bdi, preco_c_bdi, total_s_bdi, total,
        ]
        if all(valor is None or str(valor).strip() == "" for valor in valores):
            continue

        item_texto = str(item).strip() if item is not None else ""
        codigo_texto = str(codigo).strip() if codigo is not None else ""
        descricao_texto = str(descricao).strip() if descricao is not None else ""

        if not item_texto and not codigo_texto and not descricao_texto and coluna4:
            linhas.append({"tipo": "extenso", "texto": str(coluna4).strip()})
        elif item_texto and not codigo_texto:
            linhas.append(
                {
                    "tipo": "secao",
                    "item": item_texto,
                    "descricao": descricao_texto,
                    "extenso": str(coluna4).strip() if coluna4 else "",
                    "total": total,
                }
            )
        elif codigo_texto:
            linhas.append(
                {
                    "tipo": "item",
                    "item": item_texto,
                    "codigo": codigo_texto,
                    "descricao": descricao_texto,
                    "un": str(coluna4).strip() if coluna4 else "",
                    "qtd": quantidade,
                    "preco_s_bdi": preco_s_bdi,
                    "preco_c_bdi": preco_c_bdi,
                    "total_s_bdi": total_s_bdi,
                    "total": total,
                }
            )

    return linhas, totais


def obter_valor_total_orcamento(totais):
    for rotulo, valor in totais:
        if rotulo == "Total do Orçamento":
            return valor
    return totais[-1][1] if totais else None


def montar_texto_referencia(referencia_sinapi):
    if not referencia_sinapi:
        return "Referência:"
    if referencia_sinapi.startswith("Referência:"):
        return referencia_sinapi
    return f"Referência: {referencia_sinapi}"


def preencher_celula_conclusao(celula, valor_total, texto_extenso_total):
    from core.formatador_sinapi.comum import valor_em_extenso

    moeda = formatar_moeda(valor_total)
    extenso = texto_extenso_total or valor_em_extenso(valor_total)
    if extenso and extenso[0].islower():
        extenso = extenso[0].upper() + extenso[1:]

    limpar_celula_conclusao(celula)

    paragrafo = celula.paragraphs[0]
    paragrafo.alignment = None

    texto_intro = (
        "Portanto, para o reparo das anomalias de origem endógena (vícios construtivos) "
        "constatadas no condomínio, conclui-se ser necessário o valor atualizado de"
    )
    configurar_run(
        paragrafo.add_run(texto_intro),
        texto_intro,
        FONTE_REFERENCIA,
        FONTE_CONCLUSAO_TEXTO,
        negrito=False,
    )
    configurar_run(
        paragrafo.add_run(f" {moeda} "),
        f" {moeda} ",
        FONTE_REFERENCIA,
        FONTE_CONCLUSAO_MOEDA,
        negrito=True,
    )
    configurar_run(
        paragrafo.add_run(f"({extenso})."),
        f"({extenso}).",
        FONTE_REFERENCIA,
        FONTE_CONCLUSAO_EXTENSO,
        negrito=True,
    )


def atualizar_referencia_sinapi(documento, referencia_sinapi):
    if not referencia_sinapi:
        return

    texto_referencia = montar_texto_referencia(referencia_sinapi)
    for tabela in documento.tables:
        if len(tabela.rows) < 2:
            continue
        texto_linha = tabela.rows[1].cells[0].text.strip()
        if texto_linha.startswith("Referência:"):
            substituir_texto_celula(
                tabela.rows[1].cells[0],
                texto_referencia,
                FONTE_REFERENCIA,
                Pt(14),
                negrito=True,
                alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
            )


def atualizar_conclusao(documento, valor_total, texto_extenso_total):
    if valor_total is None:
        return

    for tabela in documento.tables:
        if not tabela.rows:
            continue
        texto = tabela.rows[0].cells[0].text.strip()
        if texto.startswith("Portanto, para o reparo"):
            preencher_celula_conclusao(
                tabela.rows[0].cells[0],
                valor_total,
                texto_extenso_total,
            )
            break


def gerar_caminho_saida_word(caminho_excel, caminho_saida=None):
    if caminho_saida:
        return caminho_saida
    raiz, _ = os.path.splitext(caminho_excel)
    return f"{raiz}.docx"


def resolver_caminho_template(caminho_template=None):
    if caminho_template:
        return caminho_template

    return caminho_modelo_word(TEMPLATE_PADRAO)


def gerar_word_modelo1(
    caminho_excel,
    caminho_template=None,
    caminho_saida=None,
    abrir_arquivo=True,
):
    caminho_template = resolver_caminho_template(caminho_template)
    caminho_saida = gerar_caminho_saida_word(caminho_excel, caminho_saida)

    if not os.path.isfile(caminho_excel):
        raise FileNotFoundError(f"Planilha Excel não encontrada: '{caminho_excel}'")
    if not os.path.isfile(caminho_template):
        raise FileNotFoundError(f"Template Word não encontrado: '{caminho_template}'")

    planilha = openpyxl.load_workbook(caminho_excel, data_only=True)
    if ABA_ORCAMENTO not in planilha.sheetnames:
        raise ValueError(f"A planilha não contém a aba '{ABA_ORCAMENTO}'.")

    ws_orcamento = planilha[ABA_ORCAMENTO]
    linhas_orcamento, totais = ler_orcamento_formatado(ws_orcamento)
    referencia_sinapi, texto_extenso_total = ler_metadados_planilha(ws_orcamento)
    valor_total_orcamento = obter_valor_total_orcamento(totais)

    shutil.copy2(caminho_template, caminho_saida)
    documento = Document(caminho_saida)

    tabela_orcamento, tabelas_totais = encontrar_tabelas_modelo1(documento)
    if tabela_orcamento is None:
        raise ValueError("Não foi possível localizar a tabela do orçamento analítico no template.")

    reconstruir_tabela_orcamento(tabela_orcamento, linhas_orcamento)

    if totais:
        for tabela_totais in tabelas_totais:
            preencher_tabela_totais(tabela_totais, totais)

    atualizar_referencia_sinapi(documento, referencia_sinapi)
    atualizar_conclusao(documento, valor_total_orcamento, texto_extenso_total)

    documento.save(caminho_saida)

    if abrir_arquivo:
        os.startfile(os.path.abspath(caminho_saida))

    return caminho_saida


if __name__ == "__main__":
    import sys

    caminho_planilha = (
        sys.argv[1]
        if len(sys.argv) > 1
        else "Planilha_Sintetica_Convertida_Modelo1.xlsx"
    )
    gerar_word_modelo1(caminho_planilha)
